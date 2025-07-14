import streamlit as st
import pandas as pd
import json
import uuid
import boto3
import os
import time
from io import BytesIO
from typing import List, Dict, Any, Optional, Tuple
import streamlit_authenticator as stauth
import traceback
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from src.logging_config import get_logger

# Initialize logger
logger = get_logger(__name__)

from dotenv import load_dotenv
load_dotenv()

# Page configuration
st.set_page_config(
    page_title="Course Video Description Extractor",
    page_icon="📚",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        text-align: center;
        margin-bottom: 2rem;
    }
    .success-message {
        color: #4caf50;
        font-weight: bold;
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #e8f5e8;
        margin: 1rem 0;
    }
    .error-message {
        color: #f44336;
        font-weight: bold;
        padding: 1rem;
        border-radius: 0.5rem;
        background-color: #ffeaea;
        margin: 1rem 0;
    }
    .info-box {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Session state initialization
if "course_id" not in st.session_state:
    st.session_state.course_id = f"course_{uuid.uuid4().hex[:8]}"
if "processing_status" not in st.session_state:
    st.session_state.processing_status = "idle"
if "extracted_data" not in st.session_state:
    st.session_state.extracted_data = None
if "s3_files" not in st.session_state:
    st.session_state.s3_files = []
if "authentication_status" not in st.session_state:
    st.session_state.authentication_status = None
if "name" not in st.session_state:
    st.session_state.name = None
if "username" not in st.session_state:
    st.session_state.username = None

# AWS Configuration
AWS_REGION = os.getenv("AWS_REGION", "us-east-2")
S3_BUCKET_NAME = os.getenv("S3_BUCKET_NAME")

def setup_authenticator():
    """Set up the streamlit authenticator."""
    try:
        # Check if authenticator is already in session state
        if 'authenticator' in st.session_state:
            return st.session_state.authenticator
        
        auth_config = st.secrets["auth"]
        
        # Build credentials dictionary
        credentials = {
            "usernames": {}
        }
        
        for username in auth_config["credentials"]["usernames"]:
            user_data = auth_config["credentials"][username]
            credentials["usernames"][username] = {
                "email": user_data["email"],
                "name": user_data["name"],
                "password": user_data["password"]
            }
        
        # Create authenticator
        authenticator = stauth.Authenticate(
            credentials,
            auth_config["cookie_name"],
            auth_config["cookie_key"],
            auth_config["cookie_expiry_days"]
        )
        
        # Store in session state to avoid recreation
        st.session_state.authenticator = authenticator
        logger.info("Authentication system initialized successfully")
        return authenticator
    except Exception as e:
        logger.error(f"Authentication setup failed: {str(e)}")
        st.error(f"Authentication setup failed: {str(e)}")
        return None

def check_authentication():
    """Check if user is authenticated and handle login/logout."""
    authenticator = setup_authenticator()
    
    if not authenticator:
        st.error("Authentication system not properly configured.")
        st.stop()
    
    # Perform authentication with error handling
    try:
        # Try different methods based on streamlit-authenticator version
        try:
            # Method 1: Try with parameters (newer versions)
            login_result = authenticator.login(
                location='main',
                key='login_form'
            )
        except TypeError:
            # Method 2: Try without parameters (older versions)
            try:
                login_result = authenticator.login('main', 'login_form')
            except TypeError:
                # Method 3: Try with just location
                login_result = authenticator.login('main')
        
        # Handle different return types from streamlit-authenticator
        if login_result is None:
            # Fallback to session state values if login returns None
            authentication_status = st.session_state.get('authentication_status')
            name = st.session_state.get('name')
            username = st.session_state.get('username')
        elif isinstance(login_result, tuple) and len(login_result) == 3:
            # Unpack tuple if returned
            name, authentication_status, username = login_result
        else:
            # Try to get from session state if unexpected return
            authentication_status = st.session_state.get('authentication_status')
            name = st.session_state.get('name')
            username = st.session_state.get('username')
        
        # Store authentication results in session state
        if authentication_status is not None:
            st.session_state.authentication_status = authentication_status
        if name is not None:
            st.session_state.name = name
        if username is not None:
            st.session_state.username = username
        
    except Exception as e:
        logger.error(f"Authentication error: {e}")
        st.error(f"Authentication error: {e}")
        st.error("Please check your streamlit secrets configuration.")
        st.stop()
    
    # Get current authentication status
    current_auth_status = st.session_state.get('authentication_status')
    
    # Check authentication status
    if current_auth_status == False:
        st.error('Username/password is incorrect')
        st.stop()
    elif current_auth_status == None:
        st.warning('Please enter your username and password')
        st.stop()
    elif current_auth_status:
        return True
    
    return False

def read_excel_from_upload(uploaded_file, sheet_name: str = "Outline") -> Optional[pd.DataFrame]:
    """Read Excel file from Streamlit upload."""
    try:
        logger.info(f"Reading Excel file: {uploaded_file.name}, sheet: {sheet_name}")
        df = pd.read_excel(uploaded_file, sheet_name=sheet_name, header=None)
        logger.info(f"Successfully read Excel file with {len(df)} rows")
        return df
    except Exception as e:
        logger.error(f"Error reading Excel file {uploaded_file.name}: {e}")
        st.error(f"Error reading Excel file: {e}")
        return None

def slice_for_lesson(df: pd.DataFrame, lesson_num: int) -> pd.DataFrame:
    """Extract data for a specific lesson."""
    try:
        tag = f"Lesson {lesson_num}"
        header_mask = df[0].astype(str).str.contains(tag, case=False, na=False, regex=False)
        
        if not header_mask.any():
            raise ValueError(f"Header '{tag}' not found in Column A.")
            
        start = header_mask.idxmax()
        
        # Find the next lesson header
        next_header = df[0].astype(str).str.match(r"Lesson\s+\d+", na=False)
        try:
            end = next_header[next_header & (next_header.index > start)].idxmax() - 1
        except ValueError:
            end = len(df) - 1
            
        result_df = df.loc[start + 1: end]
        logger.info(f"Successfully sliced data for Lesson {lesson_num}: {len(result_df)} rows")
        return result_df
        
    except Exception as e:
        logger.error(f"Error slicing data for Lesson {lesson_num}: {e}")
        st.error(f"Error slicing data for Lesson {lesson_num}: {e}")
        return pd.DataFrame()

def extract_video_data(lesson_df: pd.DataFrame) -> List[Tuple[str, str]]:
    """Extract video data (title and voiceover) from a lesson DataFrame."""
    try:
        video_data = []
        video_rows = lesson_df[
            lesson_df[0].astype(str).str.strip().str.lower().str.startswith("video")
        ]
        
        if video_rows.empty:
            return video_data
            
        for _, row in video_rows.iterrows():
            title = str(row[1]).strip() if pd.notna(row[1]) else "Untitled Video"
            voiceover = str(row[3]).strip() if pd.notna(row[3]) else ""
            
            # Skip if voiceover is empty or invalid
            if not voiceover or voiceover.lower() in {"nan", ""}:
                continue
                
            video_data.append((title, voiceover))
        
        logger.info(f"Extracted {len(video_data)} videos from lesson data")
        return video_data
        
    except Exception as e:
        logger.error(f"Error extracting video data: {e}")
        st.error(f"Error extracting video data: {e}")
        return []

def find_lesson_count(df: pd.DataFrame) -> int:
    """Find the number of lessons in the DataFrame."""
    try:
        lesson_pattern = df[0].astype(str).str.match(r"Lesson\s+\d+", na=False)
        count = lesson_pattern.sum()
        logger.info(f"Found {count} lessons in dataframe")
        return count
    except Exception as e:
        logger.error(f"Error counting lessons: {e}")
        st.error(f"Error counting lessons: {e}")
        return 0

def upload_to_s3(file_content: bytes, s3_key: str, content_type: str = "application/vnd.openxmlformats-officedocument.wordprocessingml.document") -> bool:
    """Upload file content to S3."""
    try:
        logger.info(f"Uploading file to S3: {s3_key}")
        s3_client = boto3.client('s3', region_name=AWS_REGION)
        s3_client.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=s3_key,
            Body=file_content,
            ContentType=content_type
        )
        logger.info(f"Successfully uploaded file to S3: {s3_key}")
        return True
    except Exception as e:
        logger.error(f"S3 upload failed for {s3_key}: {e}")
        st.error(f"S3 upload failed: {e}")
        return False

def get_s3_url(s3_key: str) -> str:
    """Get public S3 URL for a key."""
    return f"https://{S3_BUCKET_NAME}.s3.{AWS_REGION}.amazonaws.com/{s3_key}"

def create_video_docx(lesson_num: int, video_num: int, video_title: str, voiceover: str, course_id: str) -> BytesIO:
    """Create a DOCX file for a single video description."""
    try:
        logger.info(f"Creating DOCX for Lesson {lesson_num}, Video {video_num}: {video_title}")
        # Create a new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Lesson {lesson_num} - Video {video_num}: {video_title}', 0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add line break
        doc.add_paragraph()
        
        # Video description header
        desc_header = doc.add_heading('Video Description:', level=1)
        
        # Video description content
        desc_content = doc.add_paragraph(voiceover)
        desc_content.style = 'Normal'
        
        # Save to BytesIO buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        logger.info(f"Successfully created DOCX for Lesson {lesson_num}, Video {video_num}")
        return buffer
        
    except Exception as e:
        logger.error(f"Error creating DOCX file for Lesson {lesson_num}, Video {video_num}: {e}")
        st.error(f"Error creating DOCX file: {e}")
        raise



def process_course_outline(df: pd.DataFrame, course_id: str) -> Dict[str, Any]:
    """Process the course outline and extract all video data."""
    try:
        logger.info(f"Starting course outline processing for course: {course_id}")
        # Find number of lessons
        num_lessons = find_lesson_count(df)
        logger.info(f"Found {num_lessons} lessons in the course outline")
        
        if num_lessons == 0:
            logger.warning("No lessons found in the Excel file")
            return {
                "status": "error",
                "message": "No lessons found in the Excel file"
            }
        
        all_video_data = []
        s3_files = []
        
        # Process each lesson and create individual video files
        for lesson_num in range(1, num_lessons + 1):
            lesson_df = slice_for_lesson(df, lesson_num)
            video_data = extract_video_data(lesson_df)
            
            if video_data:
                logger.info(f"Processing lesson {lesson_num} with {len(video_data)} videos")
                # Create individual DOCX file for each video
                for idx, (title, voiceover) in enumerate(video_data, 1):
                    # Track video data for summary
                    all_video_data.append({
                        "lesson": lesson_num,
                        "video": idx,
                        "title": title,
                        "voiceover": voiceover
                    })
                    
                    # Create DOCX file for this individual video
                    docx_buffer = create_video_docx(lesson_num, idx, title, voiceover, course_id)
                    # Clean filename by removing special characters
                    clean_title = ''.join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
                    clean_title = clean_title.replace(' ', '_')[:50]  # Limit length and replace spaces
                    s3_key = f"video_descriptions/{course_id}/lesson_{lesson_num}_video_{idx}_{clean_title}.docx"
                    
                    if upload_to_s3(docx_buffer.getvalue(), s3_key):
                        s3_files.append({
                            "lesson": lesson_num,
                            "video": idx,
                            "title": title,
                            "s3_key": s3_key,
                            "s3_url": get_s3_url(s3_key),
                            "file_type": "docx"
                        })
                    
                    # Clean up buffer
                    docx_buffer.close()
        
        # Create course metadata for return
        course_summary = {
            "course_id": course_id,
            "total_lessons": num_lessons,
            "total_videos": len(all_video_data),
            "extraction_date": time.strftime("%Y-%m-%d %H:%M:%S"),
            "s3_files": s3_files
        }
        
        logger.info(f"Course processing completed successfully: {len(all_video_data)} videos processed, {len(s3_files)} files uploaded")
        return {
            "status": "success",
            "message": f"Successfully extracted {len(all_video_data)} video descriptions and created individual DOCX files",
            "total_lessons": num_lessons,
            "total_videos": len(all_video_data),
            "s3_files": s3_files,
            "course_data": course_summary
        }
        
    except Exception as e:
        logger.error(f"Course processing failed for {course_id}: {str(e)}")
        return {
            "status": "error",
            "message": f"Processing failed: {str(e)}"
        }

def main():
    """Main Streamlit application."""
    # Check authentication
    if not check_authentication():
        return
    
    # Log user session
    user_name = st.session_state.get('name', 'Unknown')
    logger.info(f"User session started: {user_name}")
    
    # Header
    st.markdown('<div class="main-header">📚 Course Video Description Extractor</div>', unsafe_allow_html=True)
    
    # Sidebar with user info
    with st.sidebar:
        st.write(f"👤 Welcome, {st.session_state.get('name', 'User')}")
        
        # Logout button
        authenticator = setup_authenticator()
        if authenticator:
            authenticator.logout('Logout', 'sidebar', key='logout_button')
        
        st.divider()
        
        # Course ID
        st.write(f"🆔 Course ID: `{st.session_state.course_id}`")
        
        if st.button("🔄 Generate New Course ID"):
            st.session_state.course_id = f"course_{uuid.uuid4().hex[:8]}"
            st.rerun()
    
    # Main content
    st.subheader("📤 Upload Course Outline")
    
    # File upload
    uploaded_file = st.file_uploader(
        "Choose an Excel file (.xlsx)",
        type=['xlsx'],
        help="Upload your course outline Excel file with video descriptions"
    )
    
    if uploaded_file is not None:
        # Display file info
        st.info(f"📁 File uploaded: {uploaded_file.name} ({uploaded_file.size} bytes)")
        
        # Sheet name input
        sheet_name = st.text_input("Sheet Name", value="Outline")
        
        # Process button
        if st.button("🚀 Extract Video Descriptions", type="primary"):
            if not S3_BUCKET_NAME:
                logger.error("S3_BUCKET_NAME environment variable not set")
                st.error("❌ S3_BUCKET_NAME environment variable is not set")
                return
            
            logger.info(f"Starting video description extraction for file: {uploaded_file.name}")
            with st.spinner("Processing course outline..."):
                try:
                    # Read Excel file
                    df = read_excel_from_upload(uploaded_file, sheet_name)
                    
                    if df is None:
                        st.error("❌ Failed to read Excel file")
                        return
                    
                    # Process the course outline
                    result = process_course_outline(df, st.session_state.course_id)
                    
                    if result["status"] == "success":
                        st.session_state.extracted_data = result
                        st.session_state.s3_files = result["s3_files"]
                        st.session_state.processing_status = "completed"
                        
                        # Success message
                        st.markdown(f"""
                        <div class="success-message">
                            ✅ Successfully extracted video descriptions!<br>
                            📊 Total Lessons: {result['total_lessons']}<br>
                            🎥 Total Videos: {result['total_videos']}<br>
                            📄 DOCX Files created: {len(result['s3_files'])}
                        </div>
                        """, unsafe_allow_html=True)
                        
                    else:
                        st.error(f"❌ {result['message']}")
                        
                except Exception as e:
                    st.error(f"❌ Error processing file: {str(e)}")
                    st.code(traceback.format_exc())
    
    # Display results if available
    if st.session_state.processing_status == "completed" and st.session_state.extracted_data:
        st.divider()
        st.subheader("📊 Extraction Results")
        
        result = st.session_state.extracted_data
        
        # Display summary
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Total Lessons", result['total_lessons'])
        with col2:
            st.metric("Total Videos", result['total_videos'])
        with col3:
            st.metric("S3 Files", len(result['s3_files']))
        
        # Display S3 files
        st.subheader("📄 Generated DOCX Files")
        
        for file_info in result['s3_files']:
            st.write(f"📝 **Lesson {file_info['lesson']} - Video {file_info['video']}**: {file_info['title']} - [Download DOCX]({file_info['s3_url']})")
        
        # Display sample data
        if st.checkbox("👀 Show Sample Data"):
            st.json(result['course_data'])
    
    # Instructions
    st.divider()
    st.subheader("📋 Instructions")
    st.markdown("""
    1. **Upload** your course outline Excel file (.xlsx format)
    2. **Specify** the sheet name (default: "Outline")
    3. **Click** "Extract Video Descriptions" to process the file
    4. **Download** individual DOCX files from S3 links
    
    **Excel Format Requirements:**
    - Column A: Should contain "Lesson 1", "Lesson 2", etc. headers
    - Under each lesson, rows starting with "Video" in Column A
    - Column B: Video titles
    - Column D: Voiceover scripts
    
    **Output Format:**
    - Individual DOCX files saved to S3 bucket
    - One DOCX file per video
    - Clean video descriptions without metadata
    - Ready for sharing and editing
    """)

if __name__ == "__main__":
    main() 